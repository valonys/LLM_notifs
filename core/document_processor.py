import os
import logging
from typing import List, Dict, Any, Optional, Union
from datetime import datetime
import pandas as pd
import numpy as np
from io import BytesIO

# Document processing imports
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from langchain.schema import Document as LCDocument
from utils.config import Config
from utils.error_handler import ErrorHandler

logger = logging.getLogger(__name__)

class AdvancedDocumentProcessor:
    """Advanced document processor with multi-format support and metadata extraction"""
    
    def __init__(self, config: Config):
        self.config = config
        self.error_handler = ErrorHandler()
        
        # Supported file types and their processors
        self.processors = {
            'pdf': self._process_pdf,
            'xlsx': self._process_excel,
            'xls': self._process_excel,
            'docx': self._process_docx,
            'txt': self._process_txt,
            'csv': self._process_csv
        }
        
        # Processing statistics
        self.processing_stats = {
            'files_processed': 0,
            'total_pages': 0,
            'total_characters': 0,
            'processing_errors': 0
        }
    
    def process_files(self, uploaded_files: List[Any]) -> List[LCDocument]:
        """Process multiple uploaded files with comprehensive error handling"""
        all_documents = []
        
        for file in uploaded_files:
            try:
                documents = self._process_single_file(file)
                all_documents.extend(documents)
                self.processing_stats['files_processed'] += 1
                
            except Exception as e:
                self.processing_stats['processing_errors'] += 1
                logger.error(f"Error processing file {getattr(file, 'name', 'unknown')}: {str(e)}")
                continue
        
        logger.info(f"Processed {len(all_documents)} documents from {len(uploaded_files)} files")
        return all_documents
    
    def _process_single_file(self, file) -> List[LCDocument]:
        """Process a single file based on its type"""
        try:
            file_name = getattr(file, 'name', 'unknown_file')
            file_extension = file_name.split('.')[-1].lower() if '.' in file_name else ''
            
            if file_extension not in self.processors:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Get the appropriate processor
            processor = self.processors[file_extension]
            
            # Process the file
            documents = processor(file, file_name)
            
            # Add processing metadata
            for doc in documents:
                doc.metadata.update({
                    'file_name': file_name,
                    'file_type': file_extension,
                    'processed_at': datetime.now().isoformat(),
                    'processor_version': '2.0'
                })
            
            return documents
            
        except Exception as e:
            logger.error(f"Single file processing failed: {str(e)}")
            raise
    
    def _process_pdf(self, file, file_name: str) -> List[LCDocument]:
        """Process PDF files with enhanced metadata extraction"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available for PDF processing")
        
        documents = []
        
        try:
            reader = PdfReader(file)
            total_pages = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    
                    if text.strip():  # Only process non-empty pages
                        # Extract additional metadata
                        metadata = {
                            'source': file_name,
                            'page_number': page_num + 1,
                            'total_pages': total_pages,
                            'character_count': len(text),
                            'word_count': len(text.split()),
                            'extraction_method': 'PyPDF2'
                        }
                        
                        # Add PDF metadata if available
                        if reader.metadata:
                            metadata.update({
                                'pdf_title': reader.metadata.get('/Title', ''),
                                'pdf_author': reader.metadata.get('/Author', ''),
                                'pdf_subject': reader.metadata.get('/Subject', ''),
                                'pdf_creator': reader.metadata.get('/Creator', '')
                            })
                        
                        document = LCDocument(
                            page_content=text,
                            metadata=metadata
                        )
                        documents.append(document)
                        
                        self.processing_stats['total_characters'] += len(text)
                
                except Exception as e:
                    logger.warning(f"Failed to process page {page_num + 1} of {file_name}: {str(e)}")
                    continue
            
            self.processing_stats['total_pages'] += total_pages
            
        except Exception as e:
            logger.error(f"PDF processing failed for {file_name}: {str(e)}")
            raise
        
        return documents
    
    def _process_excel(self, file, file_name: str) -> List[LCDocument]:
        """Process Excel files with sheet-aware processing"""
        documents = []
        
        try:
            # Read all sheets
            excel_data = pd.read_excel(file, sheet_name=None, engine='openpyxl')
            
            for sheet_name, df in excel_data.items():
                try:
                    # Clean the dataframe
                    df = self._clean_dataframe(df)
                    
                    if df.empty:
                        continue
                    
                    # Convert to text representation
                    text_content = self._dataframe_to_text(df, sheet_name)
                    
                    # Create comprehensive metadata
                    metadata = {
                        'source': file_name,
                        'sheet_name': sheet_name,
                        'row_count': len(df),
                        'column_count': len(df.columns),
                        'columns': list(df.columns),
                        'data_types': df.dtypes.to_dict(),
                        'missing_values': df.isnull().sum().to_dict(),
                        'extraction_method': 'pandas'
                    }
                    
                    # Add statistical summary for numeric columns
                    numeric_columns = df.select_dtypes(include=[np.number]).columns
                    if len(numeric_columns) > 0:
                        metadata['numeric_summary'] = df[numeric_columns].describe().to_dict()
                    
                    document = LCDocument(
                        page_content=text_content,
                        metadata=metadata
                    )
                    documents.append(document)
                    
                    self.processing_stats['total_characters'] += len(text_content)
                
                except Exception as e:
                    logger.warning(f"Failed to process sheet '{sheet_name}' in {file_name}: {str(e)}")
                    continue
        
        except Exception as e:
            logger.error(f"Excel processing failed for {file_name}: {str(e)}")
            raise
        
        return documents
    
    def _process_docx(self, file, file_name: str) -> List[LCDocument]:
        """Process DOCX files with paragraph-level extraction"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX processing")
        
        documents = []
        
        try:
            doc = docx.Document(file)
            
            full_text = []
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
                    paragraph_count += 1
            
            if full_text:
                text_content = '\n'.join(full_text)
                
                metadata = {
                    'source': file_name,
                    'paragraph_count': paragraph_count,
                    'character_count': len(text_content),
                    'word_count': len(text_content.split()),
                    'extraction_method': 'python-docx'
                }
                
                # Add document properties if available
                try:
                    core_props = doc.core_properties
                    metadata.update({
                        'doc_title': core_props.title or '',
                        'doc_author': core_props.author or '',
                        'doc_subject': core_props.subject or '',
                        'doc_created': core_props.created.isoformat() if core_props.created else '',
                        'doc_modified': core_props.modified.isoformat() if core_props.modified else ''
                    })
                except Exception as e:
                    logger.warning(f"Could not extract document properties: {str(e)}")
                
                document = LCDocument(
                    page_content=text_content,
                    metadata=metadata
                )
                documents.append(document)
                
                self.processing_stats['total_characters'] += len(text_content)
        
        except Exception as e:
            logger.error(f"DOCX processing failed for {file_name}: {str(e)}")
            raise
        
        return documents
    
    def _process_txt(self, file, file_name: str) -> List[LCDocument]:
        """Process plain text files"""
        documents = []
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text_content = None
            
            for encoding in encodings:
                try:
                    file.seek(0)
                    text_content = file.read().decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise ValueError("Could not decode text file with any supported encoding")
            
            if text_content.strip():
                metadata = {
                    'source': file_name,
                    'character_count': len(text_content),
                    'word_count': len(text_content.split()),
                    'line_count': len(text_content.splitlines()),
                    'extraction_method': 'direct_read'
                }
                
                document = LCDocument(
                    page_content=text_content,
                    metadata=metadata
                )
                documents.append(document)
                
                self.processing_stats['total_characters'] += len(text_content)
        
        except Exception as e:
            logger.error(f"Text processing failed for {file_name}: {str(e)}")
            raise
        
        return documents
    
    def _process_csv(self, file, file_name: str) -> List[LCDocument]:
        """Process CSV files"""
        documents = []
        
        try:
            # Try different separators
            separators = [',', ';', '\t', '|']
            df = None
            
            for sep in separators:
                try:
                    file.seek(0)
                    df = pd.read_csv(file, sep=sep)
                    if len(df.columns) > 1:  # Successfully parsed
                        break
                except Exception:
                    continue
            
            if df is None or df.empty:
                raise ValueError("Could not parse CSV file")
            
            # Clean the dataframe
            df = self._clean_dataframe(df)
            
            # Convert to text representation
            text_content = self._dataframe_to_text(df, 'CSV Data')
            
            metadata = {
                'source': file_name,
                'row_count': len(df),
                'column_count': len(df.columns),
                'columns': list(df.columns),
                'data_types': df.dtypes.to_dict(),
                'missing_values': df.isnull().sum().to_dict(),
                'extraction_method': 'pandas_csv'
            }
            
            document = LCDocument(
                page_content=text_content,
                metadata=metadata
            )
            documents.append(document)
            
            self.processing_stats['total_characters'] += len(text_content)
        
        except Exception as e:
            logger.error(f"CSV processing failed for {file_name}: {str(e)}")
            raise
        
        return documents
    
    def _clean_dataframe(self, df: pd.DataFrame) -> pd.DataFrame:
        """Clean and prepare dataframe for processing"""
        try:
            # Remove completely empty rows and columns
            df = df.dropna(how='all').dropna(axis=1, how='all')
            
            # Clean column names
            df.columns = df.columns.astype(str).str.strip()
            
            # Replace NaN with empty string for text processing
            df = df.fillna('')
            
            return df
            
        except Exception as e:
            logger.error(f"DataFrame cleaning failed: {str(e)}")
            return df
    
    def _dataframe_to_text(self, df: pd.DataFrame, sheet_name: str) -> str:
        """Convert dataframe to structured text representation"""
        try:
            lines = [f"Sheet: {sheet_name}"]
            lines.append("=" * 50)
            
            # Add basic info
            lines.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
            lines.append("")
            
            # Add column headers
            lines.append("Columns: " + " | ".join(df.columns))
            lines.append("-" * 80)
            
            # Add sample data (first 20 rows for performance)
            sample_size = min(20, len(df))
            for idx, row in df.head(sample_size).iterrows():
                row_text = " | ".join([str(val) for val in row.values])
                lines.append(row_text)
            
            if len(df) > sample_size:
                lines.append(f"... and {len(df) - sample_size} more rows")
            
            # Add summary statistics for numeric columns
            numeric_columns = df.select_dtypes(include=[np.number]).columns
            if len(numeric_columns) > 0:
                lines.append("")
                lines.append("Numeric Summary:")
                for col in numeric_columns:
                    stats = df[col].describe()
                    lines.append(f"{col}: mean={stats['mean']:.2f}, std={stats['std']:.2f}")
            
            return "\n".join(lines)
            
        except Exception as e:
            logger.error(f"DataFrame to text conversion failed: {str(e)}")
            return df.to_string()
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get document processing statistics"""
        return self.processing_stats.copy()
    
    def reset_stats(self):
        """Reset processing statistics"""
        self.processing_stats = {
            'files_processed': 0,
            'total_pages': 0,
            'total_characters': 0,
            'processing_errors': 0
        }
